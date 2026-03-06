from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import UnstructuredImageLoader
from docx import Document as DocxReader
from pdf2image import convert_from_path
import pytesseract
from PIL import ImageOps, ImageFilter

def load_text_file(filepath):
    """
    Load a text (.text) or a markdown (.md) file.

    Args:
        filepath: Path to the file
    Return:
        A list containing a single Document
    """
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    chunks = [c.strip() for c in content.split("\n\n") if c.strip()] # Return multiple smaller Document objects instead of one big one for token efficiency

    documents = []
    for i, chunk in enumerate(chunks):
        metadata = {
            "source_path": str(filepath),
            "document_type": "note",
            "title":  filepath.stem,
            "filepath": filepath.suffix,
            "chunk_index": i
        }
        document = Document(page_content=chunk, metadata=metadata)
        documents.append(document)

    return documents
    
def load_pdf_file(filepath, ocr_engine, min_chars_per_page=150, max_pages_to_check=3):
    """
    Load a PDF (.pdf) file.

    Args:
        filepath: Path to the pdf
    Return:
        A list containing a Document objects (usually per page)
    """
    loader = PyPDFLoader(str(filepath))
    documents = loader.load()

    # Check to see if a PDF has sufficient native text
    pages_to_check = min(max_pages_to_check, len(documents))
    text_pages = 0
    required = (pages_to_check+1)//2
    is_native = False

    for i in range(pages_to_check):
        content = documents[i].page_content.strip()
        if len(content) >= min_chars_per_page:
            text_pages+=1
            if text_pages >= required:
                is_native = True
                break
    
    # If PDF, just return documents
    if is_native:
        for doc in documents:
            doc.metadata.update(
                {
                    "source_path": str(filepath),
                    "document_type": "pdf",
                    "title": filepath.stem,
                    "filetype" : filepath.suffix
                }
            )
        return documents

    # Fallback to OCR
    else:
        images = convert_from_path(str(filepath), dpi=300)
        ocr_documents=[]

        for image in enumerate(images):
            image = ImageOps.grayscale(image)
            image = image.filter(ImageFilter.SHARPEN)
             
            text = pytesseract.image_to_string(image, lang="eng").strip()

            if not text:
                continue
            
            ocr_documents.append(
                Document(page_content=text,
                         metadata={
                            "source_path": str(filepath),
                            "document_type": filepath.suffix,
                            "title": filepath.stem,
                            "filetype": filepath.suffix,
                        }
                    )
            )
            return ocr_documents

def load_word_file(filepath):
    """
    Load a Word (.docx) file.

    Args:
        filepath: Path to the word doc
    Return:
        A list containing a single Document
    """
    doc = DocxReader(filepath)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()] # Fix for empty strings
    
    documents = []
    for i, paragraph in enumerate(paragraphs):
        metadata = {
            "source_path": str(filepath),
            "document_type": "word",
            "title": filepath.stem,
            "filepath": filepath.suffix
        }
        document = Document(page_content=paragraph, metadata=metadata)
        documents.append(document)
    
    return documents

def load_image_file(filepath):
    """
    Load an image and run OCR extraction on it.
    
    Args:
        filepath: Path to image
        
    Returns:
        Extracted text embedded in a Document object 
    """
    loader = UnstructuredImageLoader(filepath)
    documents = loader.load()

    for doc in documents:
        doc.metadata.update(
            {
                "source_path": str(filepath),
                "document_type": "image",
                "title": filepath.stem,
                "filetype" : filepath.suffix
            }
        )

    return documents

def load_document(filepath):
    """
    Load a document based on its file-type.

    Args:
        filepath: Path to the document

    Returns:
        List of Document objects, or None if the filetype isn't supported
    """
    suffix = filepath.suffix.lower()

    try:
        if suffix in [".txt", ".md"]:
            return load_text_file(filepath)
        elif suffix == ".pdf":
            return load_pdf_file(filepath)
        elif suffix == ".docx":
            return load_word_file(filepath)
        elif suffix in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp", ".heic"]:
            return load_image_file(filepath)
        else:
            return None
        
    except Exception as e:
        print(f"Error loading {filepath}: {e}")
        return None

def discover_documents(data_dir):
    """
    Discover all supported documents in a directory tree.

    Args:
        data_dir: Root directory to scan

    Returns:
        List of file paths
    """
    supported_types = [".txt", ".md", ".pdf", ".docx", 
                       ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp", ".heic"]
    documents = []
    
    for filepath in data_dir.rglob("*"):
        if filepath.is_file() and filepath.suffix.lower() in supported_types:
            documents.append(filepath)
    
    return sorted(documents)